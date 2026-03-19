"""Pure unit tests for .docx format manipulation.

No mocks — creates real .docx bytes and verifies their content.
"""

from __future__ import annotations

import io

from docx import Document

from apron_tools.providers.microsoft.word.document import (
    append_content,
    build_docx,
    extract_text,
)

# ---------------------------------------------------------------------------
# build_docx
# ---------------------------------------------------------------------------


class TestBuildDocx:
    def test_empty(self) -> None:
        data = build_docx()
        doc = Document(io.BytesIO(data))
        assert len(doc.paragraphs) == 0

    def test_with_content(self) -> None:
        data = build_docx("Hello World\nSecond line")
        doc = Document(io.BytesIO(data))
        texts = [p.text for p in doc.paragraphs]
        assert "Hello World" in texts
        assert "Second line" in texts

    def test_single_line(self) -> None:
        data = build_docx("Just one paragraph")
        doc = Document(io.BytesIO(data))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "Just one paragraph"

    def test_empty_string_creates_empty_doc(self) -> None:
        data = build_docx("")
        doc = Document(io.BytesIO(data))
        assert len(doc.paragraphs) == 0


# ---------------------------------------------------------------------------
# extract_text
# ---------------------------------------------------------------------------


class TestExtractText:
    def test_basic_extraction(self) -> None:
        data = build_docx("Hello World")
        result = extract_text(data)
        assert "Hello World" in result["paragraphs"]
        assert result["tables"] == []

    def test_empty_document(self) -> None:
        data = build_docx()
        result = extract_text(data)
        assert result["paragraphs"] == []
        assert result["tables"] == []

    def test_multi_paragraph(self) -> None:
        data = build_docx("First\nSecond\nThird")
        result = extract_text(data)
        assert len(result["paragraphs"]) == 3
        assert result["paragraphs"][0] == "First"
        assert result["paragraphs"][1] == "Second"
        assert result["paragraphs"][2] == "Third"

    def test_table_extraction(self) -> None:
        """Build a document with a table and verify extraction."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"
        buf = io.BytesIO()
        doc.save(buf)

        result = extract_text(buf.getvalue())
        assert len(result["tables"]) == 1
        assert result["tables"][0][0] == ["A1", "B1"]
        assert result["tables"][0][1] == ["A2", "B2"]

    def test_paragraphs_and_tables(self) -> None:
        """Document with both paragraphs and a table."""
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Col1"
        table.cell(0, 1).text = "Col2"
        doc.add_paragraph("After table")
        buf = io.BytesIO()
        doc.save(buf)

        result = extract_text(buf.getvalue())
        assert "Before table" in result["paragraphs"]
        assert "After table" in result["paragraphs"]
        assert len(result["tables"]) == 1


# ---------------------------------------------------------------------------
# append_content
# ---------------------------------------------------------------------------


class TestAppendContent:
    def test_append_to_empty(self) -> None:
        data = build_docx()
        updated = append_content(data, "New content")
        doc = Document(io.BytesIO(updated))
        texts = [p.text for p in doc.paragraphs]
        assert "New content" in texts

    def test_append_preserves_existing(self) -> None:
        data = build_docx("Original")
        updated = append_content(data, "Appended")
        doc = Document(io.BytesIO(updated))
        texts = [p.text for p in doc.paragraphs]
        assert "Original" in texts
        assert "Appended" in texts

    def test_append_multi_line(self) -> None:
        data = build_docx("First")
        updated = append_content(data, "Second\nThird")
        doc = Document(io.BytesIO(updated))
        texts = [p.text for p in doc.paragraphs]
        assert "First" in texts
        assert "Second" in texts
        assert "Third" in texts

    def test_returns_different_bytes(self) -> None:
        original = build_docx("Start")
        updated = append_content(original, "End")
        assert original != updated
