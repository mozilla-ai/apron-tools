"""Pure synchronous functions for .docx format manipulation.

No network calls, no tokens. Takes bytes in, returns bytes out.
Independently testable by creating real .docx content in memory.
"""

from __future__ import annotations

import io

from docx import Document


def build_docx(content: str = "") -> bytes:
    """Create a new Word document with optional paragraph content.

    Args:
        content: If non-empty, each line becomes a paragraph.
            Otherwise, creates an empty document.

    Returns:
        The .docx file content as bytes.
    """
    doc = Document()

    if content:
        for line in content.splitlines():
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def extract_text(docx_bytes: bytes) -> dict:
    """Extract plain text from paragraphs and tables in a Word document.

    Args:
        docx_bytes: Raw .docx file bytes.

    Returns:
        A dict containing:
        - paragraphs: list of paragraph text strings.
        - tables: list of tables, each a list of rows, each a list of cell strings.
    """
    doc = Document(io.BytesIO(docx_bytes))

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    tables: list[list[list[str]]] = []
    for table in doc.tables:
        table_data: list[list[str]] = []
        for row in table.rows:
            seen_cells: set[int] = set()
            row_data: list[str] = []
            for cell in row.cells:
                cell_id = id(cell._element)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        tables.append(table_data)

    return {
        "paragraphs": paragraphs,
        "tables": tables,
    }


def append_content(docx_bytes: bytes, content: str) -> bytes:
    """Append paragraphs to an existing Word document.

    Args:
        docx_bytes: Raw .docx file bytes.
        content: Text to append; each line becomes a new paragraph.

    Returns:
        The updated .docx file content as bytes.
    """
    doc = Document(io.BytesIO(docx_bytes))

    for line in content.splitlines():
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
